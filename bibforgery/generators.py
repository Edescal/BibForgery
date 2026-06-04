from .tools import get_data_from_file
from .libjabbrev2 import jabbreviation2
from pathlib import Path
from enum import Enum
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup, NavigableString, Tag
from io import BytesIO
import json, os
from jinja2 import Template
from datetime import datetime
from collections import defaultdict


class CitationStyle(Enum):
    ACS = 1
    APS = 2


def generate_full_json(input: str, include_citations=False):
    filepath = Path(f"output/{input}_papers.json").resolve()
    raw_data = get_data_from_file(filepath)
    source_data = json.loads(raw_data)

    def process_entry(entry):
        eid = entry.get("eid", "")

        fulldate = entry.get("prism:coverDate", "0000-00-00")
        year = int(fulldate[:4]) if len(fulldate) >= 4 else 0
        month = int(fulldate[5:7]) if len(fulldate) >= 7 else 0
        day = int(fulldate[8:10]) if len(fulldate) >= 10 else 0

        journal = entry.get("prism:publicationName", "")
        journal_abbr = jabbreviation2(journal)

        citedby_count = int(entry.get("citedby-count", 0))
        item = {
            "id": eid,
            "authors": entry.get("author", []),
            "title": entry.get("dc:title"),
            "year": year,
            "month": month,
            "day": day,
            "journal": journal,
            "journal_abbrev": journal_abbr,
            "doi": entry.get("prism:doi", ""),
            "volume": entry.get("prism:volume", ""),
            "number": entry.get("prism:issueIdentifier", ""),
            "pages": entry.get("prism:pageRange", ""),
            "citedby-count": citedby_count,
        }
        return item

    papers = []
    for entry in source_data.get("papers", []):
        item = process_entry(entry)
        eid = item["id"]
        citedby_count = item["citedby-count"]

        if citedby_count > 0 and eid and input and include_citations:
            eid_input = Path(f"output/{input}/{input}_{eid}_citedby.json").resolve()
            if not eid_input.is_file():
                print(f"[Warn] File {eid_input} not found!")
                continue

            eid_file_data = get_data_from_file(eid_input)

            try:
                cites_data = json.loads(eid_file_data)
            except json.JSONDecodeError as e:
                print(f"Error: Failed to decode JSON. {e.msg}")
                print(f"Location: Line {e.lineno}, Column {e.colno}")
                continue

            cites = cites_data.get("papers", [])

            citations = []
            for entry_c in cites:
                item_c = process_entry(entry_c)
                citations.append(item_c)

            item["citedby-papers"] = citations

        papers.append(item)

    return papers


def add_html_run(paragraph, html, bold=False, italic=False, size=11):
    soup = BeautifulSoup(html, "html.parser")

    def walk(node, b=False, i=False, sub=False, sup=False):
        if isinstance(node, NavigableString):
            text = str(node)
            if text:
                run = paragraph.add_run(text)
                run.font.name = "Arial"
                run.font.size = Pt(size)
                run.bold = bold or b
                run.italic = italic or i
                run.font.subscript = sub
                run.font.superscript = sup
            return

        if not isinstance(node, Tag):
            return

        b = b or node.name in ("b", "strong")
        i = i or node.name in ("i", "em")
        sub = sub or node.name == "sub"
        sup = sup or node.name == "sup"

        for child in node.children:
            walk(child, b, i, sub, sup)

    for child in soup.children:
        walk(child)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    rPr = OxmlElement("w:rPr")

    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    run = OxmlElement("w:r")
    run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    run.append(t)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_paragraph_spacing(para, before=0, after=4, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = line


def add_hanging_paragraph(doc, number_text, body_runs, indent_cm=0.0, number_width_cm=0.8):
    """
    Párrafo con numeración manual y sangría francesa (hanging indent).
    number_text : str  — p.ej. "[47]"
    body_runs   : list of (text, bold, italic, font_size_pt)
    """
    para = doc.add_paragraph()
    set_paragraph_spacing(para, before=0, after=6)

    left = Cm(indent_cm + number_width_cm)
    hanging = Cm(number_width_cm)

    pf = para.paragraph_format
    pf.left_indent = left
    pf.first_line_indent = -hanging

    # número
    run_num = para.add_run(number_text + "\t")
    run_num.font.name = "Arial"
    run_num.font.size = Pt(10)
    run_num.font.bold = False

    # set tab stop at number_width_cm so text aligns after the number
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "left")

    # convert cm to twips (1 cm ≈ 567 twips)
    tab.set(qn("w:pos"), str(int(number_width_cm * 567)))
    tabs.append(tab)
    pPr.append(tabs)

    for text, bold, italic, size, is_html in body_runs:
        if is_html:
            add_html_run(para, text, bold, italic, size)
        else:
            r = para.add_run(text)
            r.font.name = "Arial"
            r.font.size = Pt(size or 11)
            r.font.bold = bold
            r.font.italic = italic

    return para


def add_citation_formatted(
    doc,
    entry,
    global_index,
    abbreviated=True,
    extra_indent=False,
    citation_style: CitationStyle = CitationStyle.ACS,
):
    indent_cm = 0.8 if extra_indent else 0.0

    authors_raw = entry.get("author", [])
    title = entry.get("xml:title", entry.get("dc:title", ""))
    journal = entry.get("prism:publicationName", "")
    volume = entry.get("prism:volume", "")
    issue = entry.get("prism:issueIdentifier", "")
    pages = entry.get("prism:pageRange", "") or entry.get("page-range", "")
    year = (entry.get("prism:coverDate", "") or "")[:4]
    doi = entry.get("prism:doi", "")

    authors = ""
    number_text = f"[{global_index}]"
    journal = jabbreviation2(journal) if abbreviated else journal

    # construir runs: (text, bold, italic, pt)
    runs = []
    FONT_SIZE = 11
    if CitationStyle(citation_style) == CitationStyle.ACS:
        for i, author in enumerate(authors_raw):
            if i == len(authors_raw) - 1:
                authors += f"{author.get('surname')} {author.get('initials')}"
                continue
            authors += f"{author.get('surname')} {author.get('initials')}"
            if not i == len(authors_raw) - 1:
                authors += "; "

        if authors:
            runs.append((authors + " ", False, False, FONT_SIZE, False))
        if title:
            runs.append((title + ". ", False, False, FONT_SIZE, True))
        if journal:
            runs.append((journal, False, True, FONT_SIZE, False))
        if year:
            runs.append((f" {year},", True, False, FONT_SIZE, False))
        if volume:
            runs.append((f" {volume}", False, True, FONT_SIZE, False))
        if issue:
            runs.append((f"({issue}),", False, False, FONT_SIZE, False))
        if pages:
            runs.append((f"{pages}.", False, False, FONT_SIZE, False))

    elif CitationStyle(citation_style) == CitationStyle.APS:
        for i, author in enumerate(authors_raw):
            if i == len(authors_raw) - 1:
                authors += f"and {author.get('initials')} {author.get('surname')}."
                continue

            authors += f"{author.get('initials')} {author.get('surname')}"
            if len(authors_raw) > 1 and i == len(authors_raw) - 2:
                authors += " "
            elif not i == len(authors_raw) - 1:
                authors += ", "

        if authors:
            runs.append((authors + " ", False, False, FONT_SIZE, False))
        if title:
            runs.append((title + ". ", False, False, FONT_SIZE, True))
        if journal:
            runs.append((journal, False, True, FONT_SIZE, False))
        if volume:
            runs.append((f" {volume}", True, False, FONT_SIZE, False))
        if issue:
            runs.append((f"({issue})", False, False, FONT_SIZE, False))
        if pages:
            runs.append((f", {pages}", False, False, FONT_SIZE, False))
        if year:
            runs.append((f" ({year}).", False, False, FONT_SIZE, False))

    para = add_hanging_paragraph(doc, number_text, runs, indent_cm=indent_cm)

    if doi:
        r_prefix = para.add_run(" DOI: ")
        r_prefix.font.name = "Arial"
        r_prefix.font.size = Pt(10)

        add_hyperlink(para, doi, f"https://doi.org/{doi}")


def generate_docx(name, include_citations=False, citation_style=CitationStyle.ACS) -> None:
    filepath = Path(f"output/{name}_papers.json").resolve()
    raw_data = get_data_from_file(filepath)
    data = json.loads(raw_data)

    entries = data.get("papers", [])

    def get_year(e):
        d = e.get("prism:coverDate", "0000")
        try:
            return int(d[:4])
        except (ValueError, TypeError):
            return 0

    sorted_entries = sorted(entries, key=get_year, reverse=True)
    total = len(sorted_entries)

    doc = Document()

    # estilos base
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # márgenes y tamaño
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # título
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(6)
    r = title_para.add_run("Publications")
    r.font.bold = True
    r.font.size = Pt(12)
    r.font.name = "Arial"

    for i, entry in enumerate(sorted_entries):
        global_index = total - i
        cited_count = int(entry.get("citedby-count", 0))
        eid = entry.get("eid", "")
        short_title = entry.get("dc:title", "")[:55]
        print(f" [{global_index:>3}/{total}] {short_title}... (cited by {cited_count} papers)")

        add_citation_formatted(doc, entry, global_index, abbreviated=True, citation_style=citation_style)

        if cited_count > 0 and eid and name and include_citations:
            eid_input = Path(f"output/{name}/{name}_{eid}_citedby.json").resolve()
            if not eid_input.is_file():
                print(f"[Warn] File {eid_input} not found!")
                continue

            eid_file_data = get_data_from_file(eid_input)

            try:
                cites_data = json.loads(eid_file_data)
            except json.JSONDecodeError as e:
                print(f"Error: Failed to decode JSON. {e.msg}")
                print(f"Location: Line {e.lineno}, Column {e.colno}")
                continue

            cites = cites_data.get("papers", [])

            label_para = doc.add_paragraph()
            label_para.paragraph_format.left_indent = Cm(0.8)
            label_para.paragraph_format.space_before = Pt(3)
            label_para.paragraph_format.space_after = Pt(2)
            r = label_para.add_run(f"Cited by ({len(cites)}):")
            r.font.bold = True
            r.font.size = Pt(11)
            r.font.name = "Arial"

            for j, cite_entry in enumerate(cites, start=1):
                add_citation_formatted(
                    doc, cite_entry, j, abbreviated=True, extra_indent=True, citation_style=citation_style
                )

        # separador visual entre entradas principales
        # sep = doc.add_paragraph()
        # sep.paragraph_format.space_before = Pt(0)
        # sep.paragraph_format.space_after = Pt(0)

    buffer = BytesIO()
    doc.save(buffer)

    return buffer.getvalue()


def generate_pdf(name, include_citations=False, citation_style=CitationStyle.ACS) -> bytes | None:
    from weasyprint import HTML, CSS

    filepath = Path(f"output/{name}_papers.json").resolve()
    raw_data = get_data_from_file(filepath)
    data = json.loads(raw_data)

    entries = data.get("papers", [])

    def get_year(e):
        d = e.get("prism:coverDate", "0000")
        try:
            return int(d[:4])
        except (ValueError, TypeError):
            return 0

    sorted_entries = sorted(entries, key=get_year, reverse=True)
    total = len(sorted_entries)

    # ————
    grouped = defaultdict(list)
    for i, entry in enumerate(sorted_entries):
        global_index = total - i

        year = (entry.get("prism:coverDate", "") or "")[:4] or "Sin Año"

        citation_html = process_single_entry_as_html(entry, global_index, citation_style)
        grouped[year].append(citation_html)

        cited_count = int(entry.get("citedby-count", 0))
        eid = entry.get("eid", "")

        if not include_citations or cited_count <= 0 or not eid:
            continue

        eid_input = Path(f"output/{name}/{name}_{eid}_citedby.json").resolve()

        if not eid_input.is_file():
            print(f"[Warn] File {eid_input} not found!")
            continue

        try:
            cites_data = json.loads(get_data_from_file(eid_input))
        except json.JSONDecodeError:
            continue

        cites = cites_data.get("papers", [])

        grouped[year].append(f"""
            <div style="margin-left:30px;
                        margin-top:4px;
                        margin-bottom:4px;
                        font-weight:bold;">
                Cited by ({len(cites)}):
            </div>
            """)

        for j, cite_entry in enumerate(cites, start=1):
            grouped[year].append(process_single_entry_as_html(cite_entry, j, citation_style, extra_indent=True))

    sorted_years = sorted(grouped.keys(), reverse=True)
    grouped_data = [(y, grouped[y]) for y in sorted_years]
    # ————

    base_path = os.path.dirname(os.path.abspath(__file__))
    templates_path = os.path.join(base_path, "templates/")
    html_path = os.path.join(base_path, "templates", "index.html")
    css_path = os.path.join(base_path, "templates", "styles.css")

    with open(html_path, "r", encoding="utf-8") as html:
        plantilla_html = html.read()

    context = {
        "grouped_data": grouped_data,
        "date": datetime.now().strftime("%B %d, %Y"),
    }
    render = Template(plantilla_html).render(data=context)
    styles = CSS(filename=css_path)
    pdf_as_bytes = HTML(
        string=render,
        base_url=templates_path,
    ).write_pdf(stylesheets=[styles])

    return pdf_as_bytes


def process_single_entry_as_html(data, index: int, citation_style=CitationStyle.ACS, extra_indent=False):
    res = ""
    authors = ""
    authors_raw = data.get("author", [])

    title = data.get("xml:title", data.get("dc:title", ""))
    journal = jabbreviation2(data.get("prism:publicationName", ""))
    volume = data.get("prism:volume", "")
    issue = data.get("prism:issueIdentifier", "")
    pages = data.get("prism:pageRange", "") or data.get("page-range", "")
    year = (data.get("prism:coverDate", "") or "")[:4]
    doi = data.get("prism:doi", "")

    if CitationStyle(citation_style) == CitationStyle.ACS:
        for i, author in enumerate(authors_raw):
            if i == len(authors_raw) - 1:
                authors += f"{author.get('surname')} {author.get('initials')}"
                continue
            authors += f"{author.get('surname')} {author.get('initials')}"
            if not i == len(authors_raw) - 1:
                authors += "; "

        res = f'{authors} {title}. <span style="font-style: italic;">{journal}</span>'
        if year:
            res += f' <span style="font-weight: bold;">{year}</span>'
        if volume or issue or pages:
            res += ","
            if volume:
                res += f' <span style="font-style: italic;">{volume}</span>'
            if issue:
                res += f"({issue})"
            if pages:
                res += f", {pages}."

    elif CitationStyle(citation_style) == CitationStyle.APS:
        for i, author in enumerate(authors_raw):
            if i == len(authors_raw) - 1:
                authors += f"and {author.get('initials')} {author.get('surname')}."
                continue
            authors += f"{author.get('initials')} {author.get('surname')}"
            if not i == len(authors_raw) - 1:
                authors += ", "

        res = f'{authors} {title}. <span style="font-style: italic;">{journal}</span>'
        if volume or issue or pages:
            res += ","
            if volume:
                res += f' <span style="font-weight: bold;">{volume}</span>'
            if issue:
                res += f"({issue})"
            if pages:
                res += f", {pages},"
        if year:
            res += f" {year}."

    if doi:
        res += f' DOI: <a href="https://doi.org/{doi}" target="_blank" rel="noopener noreferrer" style="color: #0563C1; text-decoration: underline;">https://doi.org/{doi}</a>'

    # return f"""
    # <div style="display: flex; margin-bottom: 8px;">
    #     <div style="min-width: 30px; font-weight: bold;">{index}.</div>
    #     <div style="flex: 1;">{res}</div>
    # </div>
    # """

    margin = "40px" if extra_indent else "0px"
    return f"""
    <div style="
        display:flex;
        margin-bottom:8px;
        margin-left:{margin};
    ">
        <div style="min-width:30px;font-weight:bold;">
            {index}.
        </div>
        <div style="flex:1;">
            {res}
        </div>
    </div>
    """
